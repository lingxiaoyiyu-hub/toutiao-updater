# -*- coding: utf-8 -*-
"""
今日头条文章 Word (.docx) 导出器 (Docx Exporter)
=================================================
将抓取的文章标题、正文及高清配图优雅排版导出为标准 Word 文档。
"""

import os
import re
from pathlib import Path
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


class DocxExporter:
    """Word 文档生成器"""

    @staticmethod
    def export_article_to_docx(article_dict: Dict[str, Any], output_path: str) -> str:
        """导出单篇文章为带图 Word 文档"""
        doc = Document()

        # 设置页面边距
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        title = article_dict.get("title", "今日头条文章")
        author = article_dict.get("author", "今日头条创作者")
        pub_time = article_dict.get("publish_time", "")
        read_count = article_dict.get("read_count", 0)
        digg_count = article_dict.get("digg_count", 0)
        url = article_dict.get("article_url", "")
        content_md = article_dict.get("content_markdown", "")

        # 1. 主标题
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run(title)
        run_title.font.name = "Microsoft YaHei"
        run_title.font.size = Pt(18)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(26, 26, 26)

        # 2. 元数据信息栏
        p_meta = doc.add_paragraph()
        p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_text = f"作者：{author}  |  发布时间：{pub_time}  |  阅读量：{read_count}  |  点赞：{digg_count}"
        run_meta = p_meta.add_run(meta_text)
        run_meta.font.name = "Microsoft YaHei"
        run_meta.font.size = Pt(9.5)
        run_meta.font.color.rgb = RGBColor(128, 128, 128)

        # 分割线
        p_div = doc.add_paragraph()
        p_div_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/></w:pBdr>')
        p_div._p.get_or_add_pPr().append(p_div_border)

        # 3. 正文段落与图片
        lines = content_md.split("\n")
        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue

            # 忽略原 markdown 标题与 yaml
            if line_str.startswith("# ") or line_str.startswith("---") or line_str.startswith("title:") or line_str.startswith("author:"):
                continue

            # 二级小标题
            if line_str.startswith("## ") or line_str.startswith("### "):
                sub_title = line_str.lstrip("#").strip()
                p_sub = doc.add_paragraph()
                p_sub.paragraph_format.space_before = Pt(12)
                p_sub.paragraph_format.space_after = Pt(4)
                r_sub = p_sub.add_run(sub_title)
                r_sub.font.name = "Microsoft YaHei"
                r_sub.font.size = Pt(13)
                r_sub.font.bold = True
                r_sub.font.color.rgb = RGBColor(30, 80, 150)
                continue

            # 引用块
            if line_str.startswith("> "):
                quote_text = line_str.lstrip("> ").strip()
                p_q = doc.add_paragraph()
                p_q.paragraph_format.left_indent = Inches(0.3)
                r_q = p_q.add_run(quote_text)
                r_q.font.name = "Microsoft YaHei"
                r_q.font.size = Pt(10)
                r_q.font.italic = True
                r_q.font.color.rgb = RGBColor(100, 100, 100)
                continue

            # 图片
            img_match = re.search(r'!\[.*?\]\((.*?)\)', line_str)
            if img_match:
                img_path_str = img_match.group(1)
                # 检查本地图片文件
                if os.path.exists(img_path_str):
                    try:
                        p_img = doc.add_paragraph()
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_img.paragraph_format.space_before = Pt(8)
                        p_img.paragraph_format.space_after = Pt(8)
                        run_img = p_img.add_run()
                        run_img.add_picture(img_path_str, width=Inches(5.2))
                    except Exception:
                        pass
                continue

            # 普通正文段落
            p_text = doc.add_paragraph()
            p_text.paragraph_format.line_spacing = 1.25
            p_text.paragraph_format.space_after = Pt(6)
            p_text.paragraph_format.first_line_indent = Inches(0.25)
            r_text = p_text.add_run(line_str)
            r_text.font.name = "Microsoft YaHei"
            r_text.font.size = Pt(11)
            r_text.font.color.rgb = RGBColor(51, 51, 51)

        # 保存
        out_p = Path(output_path)
        out_p.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_p))
        return str(out_p)


docx_exporter = DocxExporter()
